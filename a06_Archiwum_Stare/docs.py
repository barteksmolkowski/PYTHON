from docx import Document
from docx.shared import Pt

# Tworzymy nowy dokument
doc = Document()


# Funkcja do dodawania akapitu z ustaloną wielkością czcionki
def add_paragraph(text, size=12, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p


# Tytuł
add_paragraph("UMOWA O DZIEŁO", size=14, bold=True)
add_paragraph("zawarta dnia ………………… w …………………, pomiędzy:")

# Zamawiający
add_paragraph("\nZamawiający")
add_paragraph("Imię i nazwisko: ………………………………………")
add_paragraph("Nazwa firmy: ………………………………………")

# Wykonawca
add_paragraph("\nWykonawca")
add_paragraph("Imię i nazwisko: ………………………………………")
add_paragraph("Nazwa firmy: ………………………………………")

add_paragraph("\nzwanymi dalej łącznie „Stronami”, a osobno „Stroną”.\n")

# §1 Przedmiot umowy
add_paragraph("§1. Przedmiot umowy", bold=True)
add_paragraph(
    "1. Zamawiający zleca, a Wykonawca zobowiązuje się do wykonania strony internetowej składającej się z 5 (pięciu) podstron, zgodnie z wytycznymi Zamawiającego."
)
add_paragraph(
    "2. Strona internetowa zostanie przygotowana w technologii HTML + CSS + JS, z możliwością późniejszej edycji treści przez Zamawiającego."
)

# §2 Termin realizacji
add_paragraph("\n§2. Termin realizacji", bold=True)
add_paragraph(
    "1. Wykonawca zobowiązuje się do wykonania dzieła w terminie do dnia 10 października."
)
add_paragraph(
    "2. Przekazanie dzieła nastąpi poprzez dostarczenie plików źródłowych oraz wdrożenie strony na wskazany przez Zamawiającego serwer (i adres, aby każdy normalnie mógł wejść bez problemu; w przypadku problemów odpowiedzialność leży po stronie Wykonawcy, który zobowiązuje się je natychmiast naprawić)."
)
add_paragraph(
    "3. W przypadku opóźnienia w wykonaniu dzieła z przyczyn leżących po stronie Wykonawcy, Zamawiający ma prawo naliczyć karę umowną w wysokości 3% całkowitego wynagrodzenia za każdy dzień opóźnienia."
)
add_paragraph(
    "4. Zamawiający zobowiązuje się do odbioru dzieła natychmiast po jego przekazaniu przez Wykonawcę. Brak odbioru w tym momencie uprawnia Wykonawcę do naliczenia kary umownej w wysokości 3% całkowitego wynagrodzenia za każdy dzień zwłoki w odbiorze."
)

# §3 Wynagrodzenie
add_paragraph("\n§3. Wynagrodzenie", bold=True)
add_paragraph(
    "1. Strony ustalają wynagrodzenie za wykonanie dzieła na kwotę 300 zł (bez podatków)."
)
add_paragraph("2. Wynagrodzenie zostanie wypłacone w trzech ratach:")
add_paragraph(
    "a) I rata w wysokości 33 i 1/3 % wartości umowy, płatna w terminie 3 dni od podpisania niniejszej umowy, jako zaliczka."
)
add_paragraph(
    "b) II rata w wysokości 33 i 1/3 % wartości umowy, płatna po wykonaniu 50% prac i ich akceptacji przez Zamawiającego."
)
add_paragraph(
    "c) III rata w wysokości 33 i 1/3 % wartości umowy, płatna po całkowitym wykonaniu dzieła i jego odbiorze przez Zamawiającego."
)
add_paragraph(
    "3. W przypadku opóźnienia w zapłacie którejkolwiek z rat, Zamawiający zobowiązany jest do zapłaty kary umownej w wysokości 3% całkowitego wynagrodzenia za każdy dzień opóźnienia."
)

# §4 Odbiór dzieła
add_paragraph("\n§4. Odbiór dzieła", bold=True)
add_paragraph(
    "1. Zamawiający zobowiązuje się do odbioru dzieła natychmiast po jego przekazaniu przez Wykonawcę."
)
add_paragraph(
    "2. Brak odbioru w tym momencie uprawnia Wykonawcę do naliczenia kary umownej zgodnie z §2 ust. 4."
)

# §5 Rozwiązanie umowy
add_paragraph("\n§5. Postanowienia dotyczące rozwiązania umowy", bold=True)
add_paragraph(
    "1. Umowa może być rozwiązana jednostronnie o ile zapłacono za wykonaną część pracy."
)
add_paragraph(
    "2. Rozwiązanie lub wypowiedzenie umowy jest możliwe na mocy pisemnej zgody jednej ze Stron."
)
add_paragraph(
    "3. W przypadku jednostronnego zerwania umowy przez Wykonawcę bez zgody Zamawiającego, Wykonawca zobowiązuje się do zwrotu całości otrzymanego wynagrodzenia."
)

# §6 Prawa autorskie
add_paragraph("\n§6. Prawa autorskie i materiały", bold=True)
add_paragraph(
    "1. Wszelkie materiały przesyłane Wykonawcy przez Zamawiającego pozostają własnością Zamawiającego."
)
add_paragraph(
    "2. Wykonawca nie ma prawa ich kopiować ani przekazywać osobom trzecim. W przypadku naruszenia tego zakazu, Wykonawca zobowiązany jest do natychmiastowego zwrotu wszystkich materiałów i poniesienia kary finansowej w wysokości całości wynagrodzenia."
)
add_paragraph(
    "3. Naruszenie prawa do materiałów może skutkować skierowaniem sprawy do sądu."
)
add_paragraph(
    "4. Wykonawca nie ma prawa do umieszczania informacji o autorstwie na stronie internetowej."
)

# §7 Postanowienia końcowe
add_paragraph("\n§7. Postanowienia końcowe", bold=True)
add_paragraph(
    "1. W sprawach nieuregulowanych niniejszą umową zastosowanie mają przepisy Kodeksu cywilnego."
)
add_paragraph(
    "2. Wszelkie zmiany umowy wymagają formy pisemnej pod rygorem nieważności."
)
add_paragraph(
    "3. Umowę sporządzono w dwóch jednobrzmiących egzemplarzach, po jednym dla każdej ze Stron."
)
add_paragraph(
    "4. Strony zobowiązują się do zachowania poufności co do treści niniejszej umowy. Zabrania się kopiowania, rozpowszechniania lub udostępniania jej osobom trzecim bez pisemnej zgody drugiej Strony, z wyjątkiem sytuacji wymaganych przez przepisy prawa."
)

# Podpisy
add_paragraph("\n……………………………………")
add_paragraph("(Zamawiający)")
add_paragraph("\n……………………………………")
add_paragraph("(Wykonawca)")

# Zapis dokumentu
doc.save("Umowa_o_dzielo.docx")

print("Dokument 'Umowa_o_dzielo.docx' został wygenerowany!")
